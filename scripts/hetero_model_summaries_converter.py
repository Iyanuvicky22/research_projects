from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT_DOCX = Path('../data/ecoms_results/final_results/heteroskedasticity/adrl_hetero_model_summaries.docx')
OUTPUT_DOCX = Path('../data/ecoms_results/final_results/heteroskedasticity/adrl_hetero_model_summaries_structured_tables.docx')

NUMERIC_RE = re.compile(r'^[+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:e[+-]?\d+)?$', re.I)

HEADER_LABELS = [
    'Dep. Variable', 'R-squared', 'Model', 'Adj. R-squared', 'Method', 'F-statistic',
    'Date', 'Prob (F-statistic)', 'Time', 'Log-Likelihood', 'No. Observations', 'AIC',
    'Df Residuals', 'BIC', 'Df Model', 'Covariance Type'
]
DIAG_LABELS = [
    'Omnibus', 'Durbin-Watson', 'Prob(Omnibus)', 'Jarque-Bera (JB)',
    'Skew', 'Prob(JB)', 'Kurtosis', 'Cond. No.'
]

def extract_key_values(line: str, labels):
    # Finds known statsmodels labels, then takes the text after each label up to the next label.
    # Most labels end with ':', but statsmodels prints 'Cond. No.' without a colon.
    parts = []
    for lbl in labels:
        suffix = r':?' if lbl == 'Cond. No.' else r':'
        parts.append(r'(?P<' + re.sub(r'\W+', '_', lbl).strip('_') + r'>' + re.escape(lbl) + suffix + r')')
    pattern = re.compile('|'.join(parts))
    matches = list(pattern.finditer(line))
    pairs = []
    for i, m in enumerate(matches):
        raw = m.group(0)
        key = raw[:-1] if raw.endswith(':') else raw.strip()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(line)
        value = line[m.end():end].strip()
        pairs.append((key, value))
    return pairs

def read_docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def split_reports(text: str):
    parts = re.split(r'(?=Model Report: )', text)
    return [p.strip() for p in parts if p.strip().startswith('Model Report:')]

def parse_report(block: str):
    lines = [ln.rstrip() for ln in block.splitlines()]
    country = lines[0].replace('Model Report:', '').strip()

    meta = []
    coef_rows = []
    diag = []
    notes = []

    # coefficients table lies between the long dashed rules around the var rows
    coef_header_idx = None
    for i, ln in enumerate(lines):
        if re.search(r'coef\s+std err\s+t\s+P>\|t\|', ln):
            coef_header_idx = i
            break

    if coef_header_idx is not None:
        # metadata: lines before coefficient header, excluding title/rules/OLS heading
        for ln in lines[1:coef_header_idx]:
            if not ln.strip() or set(ln.strip()) <= {'=', '-'} or 'OLS Regression Results' in ln:
                continue
            meta.extend(extract_key_values(ln, HEADER_LABELS))

        # coefficient rows start after dashed line following coef header and continue until next =====
        start = coef_header_idx + 1
        while start < len(lines) and set(lines[start].strip()) <= {'-', '='}:
            start += 1
        for ln in lines[start:]:
            s = ln.strip()
            if not s:
                continue
            if set(s) <= {'='}:
                break
            # varnames have no spaces in this output, so normal split is reliable
            toks = s.split()
            if len(toks) >= 7 and NUMERIC_RE.match(toks[1].replace('−','-')):
                var = toks[0]
                coef, std_err, t_val, p_val, ci_low, ci_high = toks[1:7]
                coef_rows.append([var, coef, std_err, t_val, p_val, ci_low, ci_high])

        # diagnostics between end of coef section and notes
        after_coef = False
        in_notes = False
        for ln in lines[start:]:
            s = ln.strip()
            if not s:
                continue
            if s.startswith('Notes:'):
                in_notes = True
                continue
            if in_notes:
                notes.append(s)
                continue
            if set(s) <= {'='}:
                after_coef = True
                continue
            if after_coef:
                diag.extend(extract_key_values(s, DIAG_LABELS))
    return country, meta, coef_rows, diag, notes

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=8)
        set_cell_shading(hdr[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_key_value_table(doc, pairs, title=None):
    rows = []
    for i in range(0, len(pairs), 2):
        left = pairs[i]
        right = pairs[i + 1] if i + 1 < len(pairs) else ('', '')
        rows.append([left[0], left[1], right[0], right[1]])
    return add_table(doc, ['Metric', 'Value', 'Metric', 'Value'], rows, widths=[1.55, 1.3, 1.55, 1.3])

def build_doc(reports, output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('ARDL Heteroskedasticity-Robust Model Summaries')
    r.bold = True
    r.font.size = Pt(16)

    intro = doc.add_paragraph('Structured tables converted from statsmodels OLS summaries. Coefficient standard errors are HAC robust where reported in the source output.')
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (country, meta, coefs, diag, notes) in enumerate(reports):
        if idx > 0:
            doc.add_page_break()
        h = doc.add_heading(f'Model Report: {country}', level=1)

        doc.add_heading('Model Fit Summary', level=2)
        # 2-column key/value table by default
        add_key_value_table(doc, meta)

        doc.add_heading('Coefficient Estimates', level=2)
        add_table(doc, ['Variable', 'Coef.', 'Std. Err.', 't', 'P>|t|', '95% CI Lower', '95% CI Upper'], coefs,
                  widths=[1.7, .9, .9, .8, .8, 1.0, 1.0])

        doc.add_heading('Diagnostics', level=2)
        add_key_value_table(doc, diag)

        if notes:
            doc.add_heading('Notes', level=2)
            for n in notes:
                p = doc.add_paragraph(style=None)
                p.paragraph_format.left_indent = Inches(0.15)
                p.add_run(n).font.size = Pt(8)

    doc.save(output_path)

if __name__ == '__main__':
    text = read_docx_text(INPUT_DOCX)
    reports = [parse_report(b) for b in split_reports(text)]
    build_doc(reports, OUTPUT_DOCX)
    print(f'Wrote {OUTPUT_DOCX} with {len(reports)} country reports')
