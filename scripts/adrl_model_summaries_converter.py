import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# -------------------------------------------------
# General helpers
# -------------------------------------------------
def read_docx_text(path: Path) -> str:
    """Read all paragraph text from a Word document."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def set_cell_shading(cell, fill: str):
    """Apply background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=8, font_color=None):
    """Set clean text in a cell with basic formatting."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)

    if font_color:
        run.font.color.rgb = RGBColor(*font_color)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="1F4E78", header_font_color=(255, 255, 255)):
    """Apply consistent styling to a Word table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(*header_font_color)


def add_table_from_rows(doc, headers, rows, title=None, col_widths=None):
    """Add a Word table from headers and row values."""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    table.allow_autofit = True

    # Header row
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=10, font_color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E78")

    # Data rows
    for row_values in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_text(row_cells[i], value, font_size=10)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    style_table(table)

    doc.add_paragraph()
    return table


# -------------------------------------------------
# Parsing helpers
# -------------------------------------------------
def split_country_blocks(text: str) -> list[tuple[str, str]]:
    """
    Split raw statsmodels-style report into:
    [
        ("Vietnam", "ARDL Model Results ..."),
        ("Brazil", "ARDL Model Results ..."),
        ...
    ]
    """
    pattern = r"Model Report:\s*(.+?)\n"
    matches = list(re.finditer(pattern, text))

    blocks = []

    for i, match in enumerate(matches):
        country = match.group(1).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block_text = text[start:end].strip()
        blocks.append((country, block_text))

    return blocks


def is_separator_line(line: str) -> bool:
    """Identify statsmodels separator lines made of === or ---."""
    stripped = line.strip()
    return bool(stripped) and set(stripped).issubset({"=", "-"})


def parse_summary_lines(lines: list[str]) -> list[list[str]]:
    """
    Parse the top statsmodels ARDL summary block into paired rows:

    Metric | Value | Metric | Value

    Handles lines like:
    Dep. Variable: gdp_growth   No. Observations: 30
    Model: ARDL(...)            Log Likelihood -225.149
    Date: Tue, 26 May 2026      AIC 480.297
    Time: 11:42:00              BIC 500.280
    Sample: 2                   HQIC 486.406
                                  30
    """

    summary = {}

    for idx, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            continue

        if is_separator_line(stripped):
            continue

        if "ARDL Model Results" in stripped:
            continue

        if "coef" in stripped and "std err" in stripped and "P>|z|" in stripped:
            break

        # Dep. Variable and No. Observations
        if "Dep. Variable:" in line:
            left = re.search(r"Dep\. Variable:\s*(.*?)\s{2,}", line)
            right = re.search(r"No\. Observations:\s*(\S+)", line)

            if left:
                summary["Dep. Variable"] = left.group(1).strip()
            if right:
                summary["No. Observations"] = right.group(1).strip()

        # Model and Log Likelihood
        elif "Model:" in line:
            left = re.search(r"Model:\s*(.*?)\s{2,}", line)
            right = re.search(r"Log Likelihood\s+(-?\S+)", line)

            if left:
                summary["Model"] = left.group(1).strip()
            if right:
                summary["Log Likelihood"] = right.group(1).strip()

        # Method and S.D. of innovations
        elif "Method:" in line:
            left = re.search(r"Method:\s*(.*?)\s{2,}", line)
            right = re.search(r"S\.D\. of innovations\s+(-?\S+)", line)

            if left:
                summary["Method"] = left.group(1).strip()
            if right:
                summary["S.D. of innovations"] = right.group(1).strip()

        # Date and AIC
        elif "Date:" in line:
            left = re.search(r"Date:\s*(.*?)\s{2,}", line)
            right = re.search(r"AIC\s+(-?\S+)", line)

            if left:
                summary["Date"] = left.group(1).strip()
            if right:
                summary["AIC"] = right.group(1).strip()

        # Time and BIC
        elif "Time:" in line:
            left = re.search(r"Time:\s*(.*?)\s{2,}", line)
            right = re.search(r"BIC\s+(-?\S+)", line)

            if left:
                summary["Time"] = left.group(1).strip()
            if right:
                summary["BIC"] = right.group(1).strip()

        # Sample and HQIC
        elif "Sample:" in line:
            left = re.search(r"Sample:\s*(.*?)\s{2,}", line)
            right = re.search(r"HQIC\s+(-?\S+)", line)

            if left:
                summary["Sample"] = left.group(1).strip()
            if right:
                summary["HQIC"] = right.group(1).strip()

            # Statsmodels often puts the sample end value on the next line
            if idx + 1 < len(lines):
                next_line = lines[idx + 1].strip()
                if re.fullmatch(r"\d+", next_line):
                    summary["Sample"] = f"{summary.get('Sample', '').strip()} to {next_line}"

    # Arrange as 4-column rows: Metric | Value | Metric | Value
    ordered_pairs = [
        ("Dep. Variable", "No. Observations"),
        ("Model", "Log Likelihood"),
        ("Method", "S.D. of innovations"),
        ("Date", "AIC"),
        ("Time", "BIC"),
        ("Sample", "HQIC"),
    ]

    rows = []

    for left_metric, right_metric in ordered_pairs:
        rows.append([
            left_metric,
            summary.get(left_metric, ""),
            right_metric,
            summary.get(right_metric, ""),
        ])

    return rows


def parse_coefficient_row(line: str):
    """
    Parse coefficient rows like:

    gdp_growth.L1 -0.0115 0.151 -0.076 0.940 -0.335 0.312

    Returns:
    [variable, coef, std_err, z, p_value, ci_low, ci_high]
    """

    pattern = (
        r"^(.+?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)\s+"
        r"(-?\d+(?:\.\d+)?(?:e[+-]?\d+)?)$"
    )

    match = re.match(pattern, line.strip(), flags=re.I)

    if not match:
        return None

    return [
        match.group(1).strip(),
        match.group(2).strip(),
        match.group(3).strip(),
        match.group(4).strip(),
        match.group(5).strip(),
        match.group(6).strip(),
        match.group(7).strip(),
    ]


def parse_country_block(block_text: str):
    """
    Parse one country block into:
    - summary_rows as 4 columns: Metric | Value | Metric | Value
    - coefficient_rows
    """

    lines = [line.rstrip() for line in block_text.splitlines() if line.strip()]

    summary_rows = parse_summary_lines(lines)
    coefficient_rows = []

    in_coef_section = False

    for line in lines:
        stripped = line.strip()

        if is_separator_line(stripped):
            continue

        if "coef" in stripped and "std err" in stripped and "P>|z|" in stripped:
            in_coef_section = True
            continue

        if in_coef_section:
            parsed_row = parse_coefficient_row(stripped)
            if parsed_row:
                coefficient_rows.append(parsed_row)

    return summary_rows, coefficient_rows


# -------------------------------------------------
# Word document builder
# -------------------------------------------------
def build_word_report(input_docx: Path, output_docx: Path):
    raw_text = read_docx_text(input_docx)
    country_blocks = split_country_blocks(raw_text)

    doc = Document()

    # Landscape layout
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    # Default font
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ARDL Model Report")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Formatted Country Summary Tables")
    subtitle_run.font.size = Pt(16)
    subtitle_run.italic = True

    doc.add_paragraph()

    for index, (country, block_text) in enumerate(country_blocks, start=1):
        if index > 1:
            doc.add_page_break()

        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"Model Report: {country}")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(31, 78, 121)

        summary_rows, coefficient_rows = parse_country_block(block_text)

        add_table_from_rows(
            doc=doc,
            headers=["Metric", "Value", "Metric", "Value"],
            rows=summary_rows,
            title="Model Summary",
            col_widths=[2.4, 2.4, 2.4, 2.4],
        )

        add_table_from_rows(
            doc=doc,
            headers=["Variable", "Coefficient", "Std. Error", "z", "P>|z|", "CI Lower", "CI Upper"],
            rows=coefficient_rows,
            title="Coefficient Estimates",
            col_widths=[2.2, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3],
        )

    doc.save(output_docx)
    print(f"Done. Word report created: {output_docx}")


if __name__ == "__main__":

    # ## ADRL Model Summaries Report Conversion
    # RESULT_FOLDER = Path("../data/ecoms_results/final_results/adrl")
    # INPUT_FILES = [
    #     "adrl_log_option_a_adrl_report.docx",
    #     "adrl_log_option_b_adrl_report.docx",
    #     "adrl_raw_model_adrl_report.docx"
    # ]

    # OUTPUT_FILES = [
    #     "adrl_log_a_model_report_word_tables.docx",
    #     "adrl_log_b_model_report_word_tables.docx",
    #     "adrl_raw_model_report_word_tables.docx"
    # ]
    

    # for input_file, output_file in zip(INPUT_FILES, OUTPUT_FILES):
    #     INPUT_DOCX = RESULT_FOLDER/input_file
    #     OUTPUT_DOCX = RESULT_FOLDER/output_file
        
    #     build_word_report(INPUT_DOCX, OUTPUT_DOCX)

